import os       # To store files to directory
import time     # To introduce time delay using sleep function
import random   # To randomize the use of the user agents
import requests     # To get HTML response of websites 
from bs4 import BeautifulSoup       # To perform scraping
from urllib.parse import urlparse   # To get domains of URLs for comparison
from newspaper import Article, ArticleException     # To check for articles on websites
from docx import Document   # To create word document
from docx.shared import Pt  # To format and style the doc
import re       # To perform text processing and pattern matching 

# List of user agents to rotate to avoid being blocked by google:
user_agents = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.101 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.114 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36 Edg/91.0.864.64",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36 OPR/77.0.4054.277"
]

# This function uses the urlparse function to get the domains of each URL to check whether we already have it.
def get_domain(url):
    parsed_url = urlparse(url)  # Takes a URL as input and returns a named tuple that contains various components of the URL.
    domain = parsed_url.netloc  # After parsing, netloc attribute of the ParseResult object contains the domain.
    return domain    

# Searches for the user specified topic on Google and returns a list of relevant URLs.
def get_search_results(topic, num_results=5):   # topic is user input, num of results can be customized (URLs)
    search_results = []     # Defining empty list to later store URLs
    headers = {'User-Agent': random.choice(user_agents)}    # Randomizing rotation of user agents before searching
    base_url = "https://www.google.com/search"      # We're going to search on google each time
    
    for start in range(0, num_results, 10):     # Step size of 10 to reduce # of requests in short time in case of large num of searches.
        params = {'q': topic, 'start': start}   # dictionary 'params' with two key-value pairs: q and start.
        response = requests.get(base_url, params=params, headers=headers)   # HTTP GET request sent to the specified base_url 
        # with the given params and headers. The base url, the user agent to use and the params dictionary specifying the topic
        # to search for and the start index, are all sent as well.
        
        if response.status_code == 200:    # Status code 200 is the standard response for a successful HTTP request.
            soup = BeautifulSoup(response.content, 'html.parser')   #  Create a BeautifulSoup object to parse the response content.
            links = soup.select('div.tF2Cxc a') # Uses the select method of the soup object to find and extract specific 
            # elements (in this case the anchor tag <a>).
            
            for link in links:       # Iterate over the list of URLs.
                url = link['href']   # Extract href attribute from each link.
                domain = get_domain(url)    # Call get_domain function to compare domain with previous URLs.
                # Check if the article exists on the URL before saving it.
                # If URL is not already present in the list and the article exists (check_article_existence returns True),
                # then append to list of URLs.
                if not any(domain == get_domain(existing_url) for existing_url in search_results):
                    if check_article_existence(url):  # Check if the article exists using the 'check_article_existence' function.
                        search_results.append(url)
                        # When the required number of URLs is retrieved, break loop.
                        if len(search_results) == num_results:
                            break
        
        time.sleep(1)  # Introduce a 1-second delay between each request.
    
    return search_results[:num_results] 

def is_xml_compatible(text):    # Check if the text is XML compatible (no NULL bytes or control characters).
    return all(ord(char) >= 32 and ord(char) <= 126 for char in text)

def get_first_paragraph(url, min_words=20):     # Function to retrieve first para from each site.
    irrelevant_strings = ["dont have permission", "don't have permission", "JavaScript and cookies", "<<", ">>", "Gateway"
    "403 - Forbidden", "Access to this page is forbidden", "It looks like you", "could not be decoded", "site requires", "gateway",
    "terms & conditions"]
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    paragraphs = soup.find_all('p')     # Criterion 1: the text must be inside a <p> tag.
    first_paragraph = None

    for paragraph in paragraphs:
        try:
            text = paragraph.get_text().strip()
            if not is_xml_compatible(text):     # Check if the text is XML compatible (no NULL bytes or control characters).
                continue

            if len(text.split()) >= min_words:  # Criterion 2: the text must be of a minimum word count (20 in this case).
                # Check if any irrelevant string matches any part of the paragraph:
                if irrelevant_strings is not None and any(irrelevant_str.lower() in text.lower() for irrelevant_str in irrelevant_strings):
                    continue
                first_paragraph = text
                break
        except UnicodeEncodeError:  # Skip paragraphs that are not XML compatible (UnicodeEncodeError).
                continue

    if not first_paragraph:
        # If no <p> tag with 20 or more words, try to find any body of text with at least 20 characters:
        all_text = soup.get_text()
        for text in all_text.split('\n'):
            text = text.strip()
            try:
                # Check if the text is XML compatible (no NULL bytes or control characters).
                if not is_xml_compatible(text):
                    continue

                if len(text) >= min_words:
                    # Check if any irrelevant string matches any part of the paragraph.
                    if irrelevant_strings is not None and any(irrelevant_str.lower() in text.lower() for irrelevant_str in irrelevant_strings):
                        continue
                    first_paragraph = text
                    break
            except UnicodeEncodeError:
                # Skip paragraphs that are not XML compatible (UnicodeEncodeError).
                continue

    return first_paragraph

def check_article_existence(url):    # Function using newspaper lib to check for existence of articles on the scraped URLs.
    try:        # If an exception of ArticleException occurs, except block will be executed instead.
        article = Article(url)  # Create an Article object from the newspaper lib.
        article.download()  # Fetch the web page's content.
        article.parse()     # Parse the downloaded content.
        return article.is_parsed       
    except ArticleException:
        return False

def replace_spaces(input_string, c):    # Generalized function to replace spaces in the topic with underscores to help make file name.
    return ''.join(c if char == ' ' else char for char in input_string)

def write_to_file(file_path, content):  # To write content to the file at runtime.
    with open(file_path, 'a', encoding='utf-8') as file:    # Open file in append mode using specified encoding.
        file.write(content + '\n')  # Write the content to file and add a new line character at the end.

def user_input():   # Function to take user input.
    topic = input("Enter a topic to search on Google: ")
    num_results = int(input("Enter the number of search results to fetch (default is 5): ") or 5)
    return topic, num_results

# This func is responsible for creating the file, setting its name according to the topic and storing it in the specified directory.  
def create_file(topic, file_directory):     
    file_dir = file_directory
    os.makedirs(file_dir, exist_ok=True)    # exist_ok=True parameter is set to avoid raising an error if the directory already exists. 
    
    base_filename = replace_spaces(topic, '_')      # Use the replace_spaces function and directory to make the file name.
    file_path = os.path.join(file_dir, f"{base_filename}.txt")
    count = 0

    while os.path.exists(file_path):    # If a file with the same name already exists, add an iterative number to the filename.
        count += 1
        file_path = os.path.join(file_dir, f"{base_filename}[{count}].txt")

    with open(file_path, 'w', encoding='utf-8'):    # Create the file if it doesn't exist.
        pass

    return file_path

# To remove duplication issues I was facing earlier, I introduced this function:
def check_paras(file_path):
    existing_paragraphs = set()     # Create an empty set, will use to retain the unique paras in our existing_paragraphs.
    with open(file_path, 'r', encoding='utf-8') as existing_file:   
        for line in existing_file:  # Iterate over each line in the opened file.
            if line.strip():    # Remove whitespaces and proceed if there is still content.
                existing_paragraphs.add(line.strip())   # Automatically only adds unique paras since it is adding to a set.
    return existing_paragraphs

def contains_invalid_characters(para):   # To remove those paras susceptible of not having our desired content.
    # Count the occurrences of each character in the para:
    count_less_than = para.count('<')
    count_greater_than = para.count('>')
    count_hash = para.count('#')
    count_slash = para.count('/')
    count_backslash = para.count('\\')
    count_left_bracket = para.count('[')
    count_right_bracket = para.count(']')

    # Check if any character occurs more than the allowed number of times:
    return (
        count_less_than > 1
        or count_greater_than > 1
        or count_hash > 1
        or count_slash > 2
        or count_backslash > 2
        or count_left_bracket > 2
        or count_right_bracket > 2
    )

def contains_high_digit_percentage(para):   # To remove those paras susceptible of not having our desired content.
    # Calculate the percentage of digits in the paragraph:
    total_chars = len(para)
    digit_chars = sum(char.isdigit() for char in para)

    if total_chars == 0:    # Handling division by 0
        return None
    else:
        digit_percentage = digit_chars / total_chars * 100

    return digit_percentage >= 30   # Check and return if the percentage of digits exceeds 30%

def append_to_file(paras, file_path, results):
    print("\nFiltered Search Results with Articles:")
    for i, url in enumerate(results, start=1):  # Iterate over the URL list: 'results'.
        print(f"{i}. {url}")
        if check_article_existence(url):    # If article exists then proceed with that URL.
            first_paragraph = get_first_paragraph(url)  # Retrieve that website's first paragraph.
            if first_paragraph and first_paragraph not in paras:  # Check if para not empty and is not already stored.
                cleaned_paragraph = re.sub(r'\[.*?\]|<.*?>', '', first_paragraph) # Use regex to remove square and angular brackets and their 
                # contents from the first_paragraph.
                
                if (
                    not contains_invalid_characters(cleaned_paragraph)
                    and not contains_high_digit_percentage(cleaned_paragraph)):
                    words = cleaned_paragraph.split()   # Split the cleaned_paragraph into words and store them into a list.
                    chunks = [words[i:i+15] for i in range(0, len(words), 15)]  # No more than 15 words in a line (to beautify the file content).

                    print(f"{cleaned_paragraph[:100]}...")      # Print a snippet of the cleaned first paragraph.

                    # write_to_file(file_path, 'URL ' + str(i) + ': ' + url + '\n')      # Add the corresponding URL for each para.
                    for chunk in chunks:    # Iterate over each chunk of 15 words.
                        line = " ".join(chunk)  # Join the words in each chunk with a " " in between all of them to create a line of words.
                        paras.add(line)         # Add this line to the set 'paras'.
                        write_to_file(file_path, f"{line}")     # Finally, write this formatted and clean content to file.
        write_to_file(file_path, '\n')      # Add a newline character after each para.

def create_doc_file(file_path, topic, file_directory):      # Using docx to create a word file from the created txt.
    capitalized_topic = topic.title()       # Capitalize the first letter of each word in the topic.
    
    base_filename = replace_spaces(topic, '_')  # Generate the base filename without any iterative number.
    doc_file_path = os.path.join(file_directory, f"{base_filename}.docx")  # Create the .doc file.

    count = 0
    # If a file with the same name already exists, add an iterative number to the filename.
    while os.path.exists(doc_file_path):
        count += 1
        doc_file_path = os.path.join(file_directory, f"{base_filename}[{count}].docx")

    doc = Document()    # Create a new Document object using the Document class from the docx library which will contain the content.

    # Add the capitalized topic heading:
    heading = doc.add_heading(capitalized_topic, level=1)   # Make a level 1 heading (top-level heading).
    heading.bold = True     # Embolden the heading.
    heading.alignment = 1   # 0 (Left), 1 (Center), 2 (Right)
    heading.runs[0].font.size = Pt(14)   # Set heading font size to 14.

    doc.add_paragraph()    # Add an empty paragraph after the heading to create space.

    with open(file_path, 'r', encoding='utf-8') as file:
        paragraphs = file.read().split('\n\n')  # Assuming that each paragraph is separated by two newline characters.

    for paragraph in paragraphs:
        # Replace multiple whitespaces with a single space within the paragraph using regex:
        formatted_paragraph = re.sub(r'\s+', ' ', paragraph.strip())

        doc.add_paragraph(formatted_paragraph)  # Add the 'formatted_paragraph' to the Word document using the 'add_paragraph()' method.

    doc.save(doc_file_path)     # Save the complete Word document to the specified 'doc_file_path'.
    print(f"Document file '{doc_file_path}' has been created.")

def create_folder():
    folder_path = r"C:\Users\hp\Desktop\scraped_content"      # Edit folder directory here to change location of the scraped data.

    if not os.path.exists(folder_path):     # Check if the folder already exists.
        os.makedirs(folder_path)        # If the folder doesn't exist, create it.
    return folder_path

# Function calls in required order:
def main():
    topic, num_results = user_input()
    file_directory = create_folder()
    file_path = create_file(topic, file_directory)
    results = get_search_results(topic, num_results)
    existing_paragraphs = check_paras(file_path)
    append_to_file(existing_paragraphs, file_path, results)
    create_doc_file(file_path, topic, file_directory)

if __name__ == "__main__":
    main()

